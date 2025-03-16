import tkinter as tk
import time
import ttkbootstrap as ttk
from ttkbootstrap.tableview import Tableview
from ttkbootstrap.dialogs import Messagebox
from ttkbootstrap.tooltip import ToolTip
import sqlite3
import json
from docx import Document
from docx.shared import Inches
import io
import matplotlib.pyplot as plt
from matplotlib import gridspec
import numpy as np
from datetime import datetime
from num2words import num2words

with open('data_file.json', encoding='UTF-8') as f:
    operation = json.load(f)
debit = operation['Debit']
credit = operation['Credit']

dat1, dat2, check = '', '', False

connection = sqlite3.connect("finance.db")
cur = connection.cursor()

cur.execute("""
    CREATE TABLE IF NOT EXISTS transactions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        date TEXT,
        type TEXT,
        amount REAL,
        comment TEXT
        )
""")
connection.commit()

columns = [
    {'text': '№ операции', "stretch": True},
    {'text': 'Дата операции', "stretch": True},
    {'text': 'Тип операции', "stretch": True},
    {'text': 'Сумма, ₴', "stretch": True},
    {'text': 'Описание операции', "stretch": True}
]



sroot = tk.Tk()
w_width = 360
w_height = 150
s_width = sroot.winfo_screenwidth()
s_height = sroot.winfo_screenheight()
x = (s_width / 2) - (w_width / 2)
y = (s_height / 2) - (w_height / 2)
sroot.geometry(f"{w_width}x{w_height}+{int(x)}+{int(y)}")
sroot.overrideredirect(1)

sroot.update_idletasks()


def main_window():
    def convert_number_to_words(number):
        return num2words(number, lang='ru', to='currency', currency='UAH')

    def rowadd():
        rows = []
        for result in cur.execute('SELECT * FROM transactions').fetchall():
            rows.append(result)
        return rows

    def selected(event):
        transaction_type = type_combobox.get()
        if transaction_type == "Доход":
            comment_combobox.configure(value=[item for item in credit])
        else:
            comment_combobox.configure(value=[item for item in debit])

    def my_amount():
        cur.execute("SELECT SUM(amount) FROM transactions")
        bals = cur.fetchone()
        bal_label.configure(text=f" Баланс :  {bals[0]:.2f} ₴,  ( {convert_number_to_words(bals[0])} )")

    def add_transaction():
        if type_combobox.get() and amount_entry.get():
            try:
                transaction_type = type_combobox.get()
                if transaction_type == "Доход":
                    amount = float(amount_entry.get())
                else:
                    amount = float(amount_entry.get()) * -1

                comment = comment_combobox.get()
                date = date_entry.entry.get()

                cur.execute("""
                    INSERT INTO transactions (date, type, amount, comment)
                    VALUES (?, ?, ?, ?)
                    """, (date, transaction_type, amount, comment))
                connection.commit()

                treeview.build_table_data(coldata=columns, rowdata=rowadd())
                treeview.reset_table()
                treeview.goto_last_page()
                type_combobox.set('')
                amount_entry.delete(0, 'end')
                comment_combobox.set('')
                Messagebox.show_info("Транзакция успешно добавлена!", "Успех")
            except ValueError:
                Messagebox.show_error("Сумма введена некорректно!", "Ошибка")
        else:
            Messagebox.show_info("Заполнены не все поля!", "Предупреждение")
        my_amount()

    def delete_transaction():
        selected_item = treeview.view.selection()[0]
        if selected_item:
            transaction_id = treeview.get_row(iid=selected_item).values[0]
            cur.execute("DELETE FROM transactions WHERE id=?", (transaction_id,))
            connection.commit()
            treeview.delete_row(iid=selected_item)
            Messagebox.show_info("Транзакция успешно удалена!", "Успех")
        else:
            Messagebox.show_info("Выберите транзакцию для удаления!", "Предупреждение")
        my_amount()

    def delete_all_transactions():
        confirm = Messagebox.yesno("Вы уверены, что хотите удалить все транзакции?", "Подтверждение")
        if confirm:
            cur.execute("DELETE FROM transactions")
            connection.commit()
            treeview.delete_rows()
            Messagebox.show_info("Все транзакции успешно удалены!", "Успех")
        my_amount()

    def edit_transaction():
        # selected_item = treeview.selection()
        selected_item = treeview.view.selection()[0]
        if not selected_item:
            Messagebox.show_error("Пожалуйста, выберите транзакцию в таблице, чтобы отредактировать.",
                                  "Транзакция не выбрана!")
            return

        transaction_id = treeview.get_row(iid=selected_item).values[0]
        date = date_entry.entry.get()
        transaction_type = type_combobox.get()
        amount = amount_entry.get()
        comment = comment_combobox.get()

        cur.execute("UPDATE transactions SET date = ?, type = ?, amount = ?, comment = ? WHERE id = ?",
                    (date, transaction_type, amount, comment, transaction_id))
        connection.commit()

        treeview.build_table_data(coldata=columns, rowdata=rowadd())
        treeview.reset_table()

        type_combobox.set('')
        amount_entry.delete(0, 'end')
        comment_combobox.set('')
        my_amount()

    def on_row_click(event):
        if treeview.view.selection():
            item = treeview.view.selection()[0]
            sel_row = treeview.get_row(iid=item)

            date_entry.entry.delete(0, 'end')
            date_entry.entry.insert(0, sel_row.values[1])
            type_combobox.set(sel_row.values[2])
            amount_entry.delete(0, 'end')
            amount_entry.insert(0, sel_row.values[3])
            comment_combobox.delete(0, 'end')
            comment_combobox.set(sel_row.values[4])

    def acc_balance():
        cur.execute("SELECT SUM(amount) FROM transactions WHERE type = ?", ("Доход",))
        dohod = cur.fetchone()
        cur.execute("SELECT SUM(amount) FROM transactions WHERE type = ?", ("Расход",))
        rashod = cur.fetchone()
        cur.execute("SELECT SUM(amount) FROM transactions")
        bals = cur.fetchone()
        connection.commit()

        Messagebox.show_info(f"Ваш общий баланс : {convert_number_to_words(bals[0])}"
                             f"\nДоход : {convert_number_to_words(dohod[0])}"
                             f"\nРасход : {convert_number_to_words(rashod[0])}",
                             "Общий баланс")

    def acc_balance_month():
        cur.execute(
            "SELECT SUM(amount) FROM transactions WHERE type = ? AND date BETWEEN date('now', 'start of month') AND date('now')",
            ("Доход",))
        dohod = cur.fetchone()
        cur.execute(
            "SELECT SUM(amount) FROM transactions WHERE type = ? AND date BETWEEN date('now', 'start of month') AND date('now')",
            ("Расход",))
        rashod = cur.fetchone()
        cur.execute(
            "SELECT SUM(amount) FROM transactions WHERE date BETWEEN date('now', 'start of month') AND date('now')")
        bals = cur.fetchone()
        connection.commit()

        Messagebox.show_info(f"Ваш баланс за месяц : {convert_number_to_words(bals[0])}"
                             f"\nДоход : {convert_number_to_words(dohod[0])}"
                             f"\nРасход : {convert_number_to_words(rashod[0])}",
                             "Баланс за теущий месяц")

    def acc_balance_period():
        cur.execute(
            "SELECT SUM(amount) FROM transactions WHERE type = ? AND date >= ? AND date <= ?",
            ("Доход", dat1, dat2,))
        dohod = cur.fetchone()
        cur.execute(
            "SELECT SUM(amount) FROM transactions WHERE type = ? AND date >= ? AND date <= ?",
            ("Расход", dat1, dat2,))
        rashod = cur.fetchone()
        cur.execute("SELECT SUM(amount) FROM transactions WHERE date >= ? AND date <= ?", (dat1, dat2))
        bals = cur.fetchone()
        connection.commit()

        Messagebox.show_info(f"Ваш баланс за период: {convert_number_to_words(bals[0])}"
                             f"\nДоход : {convert_number_to_words(dohod[0])}"
                             f"\nРасход : {convert_number_to_words(rashod[0])}",
                             "Баланс за выбранный период")

    def func(pct, allvals):
        absolute = int(np.round(pct / 100. * np.sum(allvals)))
        # return f"{pct:.1f}%\n({absolute:d})"
        return f"{pct:.1f}%"

    def pie_build(size, labels, shift, title_):
        doc = Document('mybuhhome.docx')
        doc.add_heading(title_, 1)
        explode = shift  # explode=explode,
        fig = plt.figure(figsize=[12, 4], frameon=True)
        spec = gridspec.GridSpec(ncols=2, nrows=1, figure=fig)
        # fig, axs = plt.subplots(1, 2, figsize=[10, 7], subplot_kw=dict(aspect="equal"), squeeze=True)
        ax1 = fig.add_subplot(spec[0, 0])
        ax1.set_title(title_, y=1.05, fontsize=10, fontweight='bold').set_color('blue')
        wedges, texts, autotexts = ax1.pie(size, autopct=lambda pct: func(pct, size), shadow=True, startangle=-45,
                                           wedgeprops={'width': 0.75, 'lw': 0.7, 'ls': '-', 'edgecolor': "k"},
                                           textprops=dict(color="w"), rotatelabels=True)
        bbox_props = dict(boxstyle="round4, pad=0.3", fc="w", ec="k", lw=0.72)
        kw = dict(arrowprops=dict(arrowstyle="<-"), bbox=bbox_props, zorder=0, va="center")
        plt.setp(autotexts, size=9, weight="bold")
        for i, p in enumerate(wedges):
            ang = (p.theta2 - p.theta1) / 2. + p.theta1
            y = np.sin(np.deg2rad(ang))
            x = np.cos(np.deg2rad(ang))
            horizontalalignment = {-1: "right", 1: "left"}[int(np.sign(x))]
            connectionstyle = f"angle,angleA=0,angleB={ang}"
            kw["arrowprops"].update({"connectionstyle": connectionstyle})
            ax1.annotate(labels[i], xy=(x, y), xytext=(1.2 * np.sign(x), 1.2 * y), fontsize=8,
                         horizontalalignment=horizontalalignment,
                         **kw)
        # ax.legend(wedges, labels, loc="center left", bbox_to_anchor=(1, 0, 0.5, 1), fontsize='small', shadow=True)
        # fig, ax1 = plt.subplots(1, 2, figsize=(7, 4), subplot_kw=dict(aspect="equal"), squeeze=True)
        ax2 = fig.add_subplot(spec[0, 1])
        bars = ax2.barh(labels, size, alpha=0.6, height=0.5, color=["#DDA0DD"], edgecolor="k", linewidth=2)
        ax2.set_title(title_, y=1.05, fontsize=10, fontweight='bold').set_color('maroon')
        ax2.spines[['right', 'top', 'bottom']].set_visible(False)
        ax2.xaxis.set_visible(False)
        ax2.yaxis.set_tick_params(labelsize=8)
        ax2.bar_label(bars, padding=5, color='blue', fontsize=8, label_type='edge', fmt='{:,.2f} ₴')
        fig.subplots_adjust(wspace=0.7, hspace=0.1)
        buf = io.BytesIO()
        plt.savefig(buf, format='png')
        buf.seek(0)
        # Добавляем диаграмму в Word-документ
        doc.add_picture(buf, width=Inches(5))
        # doc.add_paragraph(title_)
        doc.add_paragraph('                ')
        doc.save('mybuhhome.docx')
        plt.show()

    def stat_db():
        ssql = "SELECT comment, SUM(amount) * -1 FROM transactions WHERE type = ? GROUP BY comment"
        par1 = ("Расход",)
        labels = []
        size = []
        shift = []
        if check:
            ssql = "SELECT comment, SUM(amount) * -1 FROM transactions WHERE type = ? AND (date >= ? AND date <= ?) GROUP BY comment"
            par1 = ("Расход", dat1, dat2)
        for ret in cur.execute(ssql, par1).fetchall():
            labels.append(ret[0])
            size.append(ret[1])
            shift.append(0.1)
        pie_build(size, labels, shift, f'Расходы за период {dat1} - {dat2}')
        # explode = shift
        # fig, ax = plt.subplots()
        # ax.set_title('Расходы')
        # ax.pie(size, labels=labels, autopct='%1.1f%%', shadow=True, explode=explode,
        #        wedgeprops={'lw': 1, 'ls': '-', 'edgecolor': "k"}, rotatelabels=True)
        # ax.axis("equal")
        # plt.show()

    def stat_cr():
        ssql = "SELECT comment, SUM(amount) FROM transactions WHERE type = ? GROUP BY comment"
        par1 = ("Доход",)
        labels = []
        size = []
        shift = []
        if check:
            ssql = "SELECT comment, SUM(amount) FROM transactions WHERE type = ? AND (date >= ? AND date <= ?) GROUP BY comment"
            par1 = ("Доход", dat1, dat2)
        for ret in cur.execute(ssql, par1).fetchall():
            labels.append(ret[0])
            size.append(ret[1])
            shift.append(0.1)
        pie_build(size, labels, shift, f'Доходы за период {dat1} - {dat2}')
        # explode = shift
        # fig, ax = plt.subplots()
        # ax.set_title('Доходы')
        # ax.pie(size, labels=labels, autopct='%1.1f%%', shadow=True, explode=explode,
        #        wedgeprops={'lw': 1, 'ls': '-', 'edgecolor': "k"}, rotatelabels=True)
        # ax.axis("equal")
        # plt.show()

    def dinamika_tr():
        txtsql = 'Статистика транзакций:' + '\nДоходы:\n'
        nnum = 0

        if check == True:
            ssql = "SELECT comment, SUM(amount), COUNT(comment) FROM transactions WHERE type = ? AND (date >= ? AND date <= ?) GROUP BY comment"
            par1 = ("Доход", dat1, dat2)
            par2 = ("Расход", dat1, dat2)
        else:
            ssql = "SELECT comment, SUM(amount), COUNT(comment) FROM transactions WHERE type = ? GROUP BY comment"
            par1 = ("Доход",)
            par2 = ("Расход",)

        for ret in cur.execute(ssql, par1).fetchall():
            nnum += 1
            if len(str(nnum)) == 1:
                num = str(nnum) + '_'
            else:
                num = str(nnum)

            txtsql += f'{str(num).ljust(3)} {ret[0]}, {ret[1]:.2f} ₴, ({str(ret[2]).rjust(4)})\n'
        txtsql = '\n' + txtsql + 'Расходы:\n'
        nnum = 0
        for ret in cur.execute(ssql, par2).fetchall():
            nnum += 1
            if len(str(nnum)) == 1:
                num = str(nnum) + '_'
            else:
                num = str(nnum)

            txtsql += f'{str(num).ljust(3)} {ret[0]}, {ret[1]:.2f} ₴, ({str(ret[2]).rjust(4)})\n'
        Messagebox.show_info(txtsql, "Статистика...")

    def open_stat_win():
        win = ttk.Toplevel(title="Статистика", resizable=(False, False))
        win.iconphoto(False, photo2)
        win.geometry('+400+200')
        win.grab_set()
        class MyDateEntry(ttk.DateEntry):
            # override function
            def _on_date_ask(self):
                super()._on_date_ask()
                # generate the virtual event
                self.event_generate("<<DateEntrySelected>>")
            # function to return the selected date
            def get_date(self):
                return datetime.strptime(self.entry.get(), self._dateformat)
        def date_enabled():
            global check
            check = var1.get()
            # var_lab.configure(text='Состояние: ' + str(check))
            if check:
                var_lab.configure(image=on_i)
            else:
                var_lab.configure(image=off_i)
        def date_select(e):
            global dat1, dat2
            dat1 = date_entry1.entry.get()
            dat2 = date_entry2.entry.get()
        var_frame = ttk.LabelFrame(win, bootstyle="primary", text='Режим отчетов', relief='raised', border=1)
        var_frame.grid(row=0, column=0, padx=5, pady=5, sticky='nsew')
        time_frame = ttk.LabelFrame(win, bootstyle="primary", text='Выбор периода', relief='raised', border=1)
        time_frame.grid(row=1, column=0, padx=5, pady=5, sticky='nsew')
        stat_frame = ttk.LabelFrame(win, bootstyle="dark", text='Анализ данных', relief='raised', border=1)
        stat_frame.grid(row=2, column=0, padx=5, pady=5, sticky='nsew')
        var1 = ttk.BooleanVar()
        var1.set(False)
        var_box = ttk.Checkbutton(var_frame, bootstyle="danger",
                                  text='Учитывать начальную и конечную даты выборки данных', width=80,
                                  variable=var1, command=date_enabled)
        ToolTip(var_box, text="Учитываем начальную и конечную даты для выборки данных", wraplength=150,
                bootstyle=('danger', 'inverse'))
        var_box.grid(row=0, column=0, padx=25, pady=10, sticky='ns')
        # var_lab = ttk.Label(var_frame, bootstyle="danger", image=off_i, compound='right', text='Состояние: ' + str(var1.get()))
        var_lab = ttk.Label(var_frame, text='Состояние: ', image=off_i, compound='right')
        var_lab.grid(row=0, column=1, padx=25, sticky='ew')
        lab_st = ttk.PhotoImage(file='calendar_.png')
        date_label1 = ttk.Label(time_frame, bootstyle="info", image=lab_st, compound='left',
                                text="Начальная дата :")
        date_label1.grid(row=0, column=0, pady=5, padx=25)
        date_entry1 = MyDateEntry(time_frame, bootstyle="dark", firstweekday=0, dateformat='%Y-%m-%d', width=20)
        date_entry1.grid(row=0, column=1, pady=5, padx=5)
        date_entry1.bind('<<DateEntrySelected>>', date_select, add=True)
        wdat_sep = ttk.Separator(time_frame, bootstyle="secondary", orient='vertical')
        wdat_sep.grid(row=0, column=2, pady=3, padx=50, sticky='ew')
        date_label2 = ttk.Label(time_frame, bootstyle="info", image=lab_st, compound='left', text="Конечная дата :")
        date_label2.grid(row=0, column=3, pady=5, padx=5)
        date_entry2 = MyDateEntry(time_frame, bootstyle="dark", firstweekday=0, dateformat='%Y-%m-%d', width=20)
        date_entry2.grid(row=0, column=4, pady=5, padx=5)
        date_entry2.bind('<<DateEntrySelected>>', date_select, add=True)
        bal_button = ttk.Button(stat_frame, image=wb1, compound='left', text="Показать общий баланс",
                                width=30, command=acc_balance)
        ToolTip(bal_button, text="Показать текущий баланс, за весь период", wraplength=150)
        bal_button.grid(row=0, column=0, pady=5, padx=5, sticky='ew')
        balm_button = ttk.Button(stat_frame, compound='left', image=wb2, text="Текущий баланс за месяц",
                                 width=30, command=acc_balance_month)
        ToolTip(balm_button, text="Показать баланс за текущий месяц", wraplength=150)
        balm_button.grid(row=0, column=1, pady=5, padx=5, sticky='ew')
        balp_button = ttk.Button(stat_frame, compound='left', image=lab1, text="Баланс за выбранный период",
                                 width=30, command=acc_balance_period)
        ToolTip(balp_button, text="Показать баланс за выбранный период", wraplength=150)
        balp_button.grid(row=0, column=2, pady=5, padx=5, sticky='ew')
        wb_sep = ttk.Separator(stat_frame, bootstyle="secondary")
        wb_sep.grid(row=1, column=0, columnspan=3, pady=5, padx=2, sticky='ew')
        st1_button = ttk.Button(stat_frame, bootstyle="secondary", image=wb4, compound='left',
                                text="Диаграмма < Расход >", width=30, command=stat_db)
        ToolTip(st1_button, text="Построение диаграммы расходов", wraplength=150)
        st1_button.grid(row=2, column=0, pady=5, padx=5)
        st2_button = ttk.Button(stat_frame, bootstyle="secondary", image=wb5, compound='left',
                                text="Диаграмма < Доход >", width=30, command=stat_cr)
        ToolTip(st2_button, text="Построение диаграммы доходов", wraplength=150)
        st2_button.grid(row=2, column=1, pady=5, padx=5)
        st3_all_button = ttk.Button(stat_frame, bootstyle="secondary", image=wb6, compound='left',
                                    text="Динамика транзакций", width=30, command=dinamika_tr)
        ToolTip(st3_all_button, text="Показать динамику транзакций", wraplength=150)
        st3_all_button.grid(row=2, column=2, pady=5, padx=5)

    root = ttk.Window(title="Домашняя Бухгалтерия", themename='morph', resizable=(False, False))
    photo = ttk.PhotoImage(file='coin.png')
    root.iconphoto(False, photo)
    photo2 = ttk.PhotoImage(file='summary.png')
    wb1 = ttk.PhotoImage(file="lot-of-cash_.png")
    wb2 = ttk.PhotoImage(file='coins_.png')
    wb4 = ttk.PhotoImage(file='stat-down_.png')
    wb5 = ttk.PhotoImage(file='stat-up_.png')
    wb6 = ttk.PhotoImage(file='database-stats_.png')
    on_i = ttk.PhotoImage(file="on16.png")
    off_i = ttk.PhotoImage(file='off16.png')

    left_frame = ttk.LabelFrame(root, text='Операции', relief='raised', border=1)
    left_frame.grid(row=0, column=0, rowspan=11, padx=5, pady=5, sticky='nsew')
    right_frame = ttk.LabelFrame(root, text='Операции с транзакциями', relief='raised', border=1)
    right_frame.grid(row=0, column=1, pady=5, padx=5, ipady=2, sticky='new')  # columnspan=3,
    data_frame = ttk.Frame(root, border=1, height=50)
    data_frame.grid(row=1, column=1, rowspan=10, pady=5, padx=5, sticky='nsew')
    footer_frame = ttk.Frame(root, border=1, height=50)
    footer_frame.grid(row=11, column=0, columnspan=2, pady=5, padx=5, sticky='nsew')

    lab1 = ttk.PhotoImage(file='calendar_.png')
    date_label = ttk.Label(left_frame, bootstyle="info", image=lab1, compound='left', text="Дата операции :")
    date_label.grid(row=0, column=0, pady=5, padx=5)
    date_entry = ttk.DateEntry(left_frame, firstweekday=0, dateformat='%Y-%m-%d', width=12)
    ToolTip(date_entry, text="Выберите дату операции", wraplength=150)
    date_entry.grid(row=1, column=0, pady=5, padx=5)

    lab2 = ttk.PhotoImage(file='message-text_.png')
    type_label = ttk.Label(left_frame, bootstyle="info", image=lab2, compound='left', text="Тип операции :")
    type_label.grid(row=2, column=0, pady=5, padx=5)
    type_combobox = ttk.Combobox(left_frame, values=['Доход', 'Расход'], state="readonly")
    ToolTip(type_combobox, text="Выберите тип операции", wraplength=150)
    type_combobox.grid(row=3, column=0, pady=5, padx=5)
    type_combobox.bind('<<ComboboxSelected>>', selected)

    lab3 = ttk.PhotoImage(file='cash_.png')
    amount_label = ttk.Label(left_frame, bootstyle="danger", image=lab3, compound='left', text="Сумма операции, ₴")
    amount_label.grid(row=4, column=0, pady=5, padx=5)
    amount_entry = ttk.Entry(left_frame)
    ToolTip(amount_entry, text="Внесите сумму операции", wraplength=150)
    amount_entry.grid(row=5, column=0, pady=5, padx=5)

    lab4 = ttk.PhotoImage(file='notes_.png')
    comment_label = ttk.Label(left_frame, bootstyle="info", image=lab4, compound='left', text="Описание операции :")
    comment_label.grid(row=6, column=0, pady=5, padx=5)
    comment_combobox = ttk.Combobox(left_frame, values=[''], state="readonly")
    ToolTip(comment_combobox, text="Выберите описание операции", wraplength=150)
    comment_combobox.grid(row=7, column=0, pady=5, padx=5)

    but1 = ttk.PhotoImage(file='page-plus_.png')
    add_button = ttk.Button(left_frame, image=but1, compound='left', text="Добавить транзакцию", command=add_transaction)
    ToolTip(add_button, text="Добавление транзакции в базу", wraplength=150)
    add_button.grid(row=9, column=0, pady=10, padx=5, sticky='ew')

    #sep1 = ttk.Separator(left_frame, bootstyle="secondary")
    #sep1.grid(row=9, column=0, pady=5, padx=2, sticky='ew')

    but2 = ttk.PhotoImage(file='okrs_.png')
    st_button = ttk.Button(left_frame, image=but2, compound='left', text="Анализ транзакций", command=open_stat_win)
    ToolTip(st_button, text="Перейти к анализу транзакций", wraplength=150)
    st_button.grid(row=11, column=0, pady=10, padx=5, sticky='ew')

    #sep2 = ttk.Separator(left_frame, bootstyle="secondary")
    #sep2.grid(row=11, column=0, pady=5, padx=2, sticky='ew')

    but3 = ttk.PhotoImage(file='exit_.png')
    ex_button = ttk.Button(left_frame, image=but3, compound='left', text="Закончить работу", command=lambda: root.destroy())
    ToolTip(ex_button, text="Закрыть программу", wraplength=150)
    ex_button.grid(row=13, column=0, pady=10, padx=5, sticky='ew')

    #sep3 = ttk.Separator(left_frame, bootstyle="secondary")
    #sep3.grid(row=15, column=0, pady=5, padx=2, sticky='ew')

    bals = ttk.PhotoImage(file='balans.png')
    bal_label = ttk.Label(footer_frame, bootstyle="danger", image=bals, compound='left', text=" ")
    bal_label.grid(row=0, column=0, pady=5, padx=5, sticky='ew')
    my_amount()

    but4 = ttk.PhotoImage(file='page-edit_.png')
    edit_button = ttk.Button(right_frame, bootstyle="secondary", image=but4, compound='left',
                             text="Редактировать транзакцию", width=25, command=edit_transaction)
    edit_button.grid(row=0, column=0, pady=15, padx=5)

    but5 = ttk.PhotoImage(file='page-minus_.png')
    delete_button = ttk.Button(right_frame, bootstyle="secondary", image=but5, compound='left',
                               text="Удалить эту транзакцию", width=25, command=delete_transaction)
    delete_button.grid(row=0, column=1, pady=15, padx=5)

    but6 = ttk.PhotoImage(file='multiple-pages-minus_.png')
    delete_all_button = ttk.Button(right_frame, bootstyle="secondary", image=but6, compound='left',
                                   text="Удалить все транзакции", width=25, command=delete_all_transactions)
    delete_all_button.grid(row=0, column=2, pady=15, padx=5)

    #but_sep = ttk.Separator(right_frame, bootstyle="secondary")
    #but_sep.grid(row=1, column=0, columnspan=3, pady=5, padx=2, sticky='ew')
    colors = root.style.colors
    treeview = Tableview(
        master=data_frame,
        coldata=columns,
        rowdata=rowadd(),
        paginated=True,
        autoalign=False,
        autofit=True,
        pagesize=18,
        searchable=True,
        bootstyle='superhero',
        height=20,
    )
    treeview.pack(side="left", fill="both", expand=1, pady=5)
    treeview.goto_last_page()
    treeview.view.bind("<<TreeviewSelect>>", on_row_click)

    root.update_idletasks()
    s = root.geometry()
    s = s.split('+')
    s = s[0].split('x')
    width_root = int(s[0])
    height_root = int(s[1])

    w = root.winfo_screenwidth()
    h = root.winfo_screenheight()
    w = w // 2
    h = h // 2
    w = w - width_root // 2
    h = h - height_root // 2
    root.geometry('+{}+{}'.format(w, h))
    root.update_idletasks()
    root.mainloop()
    connection.close()


mframe = tk.Frame(sroot, width=360, height=150, background='black').place(x=0, y=0)
label1 = tk.Label(mframe, text="Домашняя Бухгалтерия", fg='white', bg='black',
                  font=('AurachBi', 26)).place(x=10, y=35)
label2 = tk.Label(mframe, text="загрузка ...", fg='white', bg='black', font=('Calibri', 11)).place(x=10, y=120)

for i in range(2):
    sroot.update_idletasks()
    time.sleep(2)

sroot.destroy()
main_window()
sroot.mainloop()
# fig, axs = plt.subplots(1, 2, figsize=(7, 4))
#    #ax.set_title('Расходы')
#    axs[0].pie(size, labels=labels, autopct='%1.1f%%', shadow=True, explode=explode, wedgeprops={'lw':1, 'ls':'-','edgecolor':"k"}, rotatelabels=True)
#    axs[1].bar(labels, size)
#    #axs.axis("equal")
#    plt.show()
